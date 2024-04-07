import tkinter as tk
from tkinter import ttk
from tkinter import messagebox
import json
import os
import datetime
from docx import Document
import sys
from flights import center_window
from reg_func import hash_password, check_password
from log_reg import  load_user_data, save_user_data
from ui import validate_length

# Путь к файлу иконки
icon_path = "Flight.ico"

class FlightSchedulerApp:
    def __init__(self, root, username):
        self.root = root
        self.username = username

        # Устанавливаем иконку
        self.root.wm_iconbitmap(True, icon_path)

        self.root.title("Flight Scheduler - {}".format(username))
        self.root.geometry("1400x700")

        self.mode = tk.StringVar()
        self.mode.set("Add")  # Исходный режим - добавление

        self.create_widgets()

        # Добавим загрузку ранее сохраненных данных при инициализации
        self.update_table()

        self.root.update_idletasks()
        width = self.root.winfo_width()
        height = self.root.winfo_height()
        x = (self.root.winfo_screenwidth() // 2) - (width // 2)
        y = (self.root.winfo_screenheight() // 2) - (height // 2)
        self.root.geometry('{}x{}+{}+{}'.format(width, height, x, y))

        # Устанавливаем действие при закрытии окна
        self.root.protocol("WM_DELETE_WINDOW", self.on_close_window)
        self.editing_flight = False  # Флаг для отслеживания редактирования рейса

    def on_close_window(self):
        # Действие при закрытии окна (например, завершение программы)
        self.root.destroy()
        sys.exit()

    def create_widgets(self):
        # Переключатели
        self.mode_selector_add = ttk.Radiobutton(self.root, text="Добавить рейс", variable=self.mode, value="Add", command=self.show_add_widgets)
        self.mode_selector_add.grid(row=0, column=0, padx=10, pady=10)

        self.mode_selector_remove = ttk.Radiobutton(self.root, text="Удалить рейс", variable=self.mode, value="Remove", command=self.show_remove_widgets)
        self.mode_selector_remove.grid(row=0, column=1, padx=10, pady=10)

        self.mode_selector_search = ttk.Radiobutton(self.root, text="Найти рейс", variable=self.mode, value="Search", command=lambda: (self.on_change_mode(), self.show_search_widgets()))
        self.mode_selector_search.grid(row=0, column=2, padx=10, pady=10)

        self.cost_label = tk.Label(self.root, text="Стоимость рейса (в рублях):")
        self.cost_label.grid(row=3, column=0, padx=10, pady=10, sticky="e")

        self.cost_entry = ttk.Entry(self.root, validate="key", validatecommand=(self.root.register(self.validate_cost), "%P"))
        self.cost_entry.grid(row=3, column=1, padx=10, pady=10, sticky="w")

        self.flight_number_label = tk.Label(self.root, text="Номер рейса:")
        self.flight_number_label.grid(row=4, column=0, padx=10, pady=10, sticky="e")

        self.flight_number_entry = ttk.Entry(self.root, validate="key", validatecommand=(self.root.register(self.validate_flight_number), "%P"))
        self.flight_number_entry.grid(row=4, column=1, padx=10, pady=10, sticky="w")

        # Поля ввода
        self.from_label = tk.Label(self.root, text="Пункт отправления:")
        self.from_label.grid(row=1, column=0, padx=10, pady=10, sticky="e")

        self.from_entry = ttk.Entry(self.root, validate="key", validatecommand=(self.root.register(self.alphabetic_validator), "%P"))
        self.from_entry.grid(row=1, column=1, padx=10, pady=10, sticky="w")

        self.to_label = tk.Label(self.root, text="Пункт прибытия:")
        self.to_label.grid(row=2, column=0, padx=10, pady=10, sticky="e")

        self.to_entry = ttk.Entry(self.root, validate="key", validatecommand=(self.root.register(self.alphabetic_validator), "%P"))
        self.to_entry.grid(row=2, column=1, padx=10, pady=10, sticky="w")

        # Поля ввода даты
        self.departure_date_label = tk.Label(self.root, text="Дата отправления (дд.мм.гггг):")
        self.departure_date_label.grid(row=5, column=0, padx=10, pady=10, sticky="e")

        self.departure_date_var = tk.StringVar()
        self.departure_date_var.trace_add("write",
                                          lambda *args: self.validate_date_entry(self.departure_date_var, 10, [2, 5]))

        self.departure_date_entry = ttk.Entry(self.root, textvariable=self.departure_date_var)
        self.departure_date_entry.grid(row=5, column=1, padx=10, pady=10, sticky="w")
        self.departure_date_entry.bind("<KeyRelease>", lambda event: self.move_cursor(event, self.departure_date_entry))

        # Поля ввода времени
        self.departure_time_label = tk.Label(self.root, text="Время отправления (чч:мм):")
        self.departure_time_label.grid(row=6, column=0, padx=10, pady=10, sticky="e")

        self.departure_time_var = tk.StringVar()
        self.departure_time_var.trace_add("write",
                                          lambda *args: self.validate_time_entry(self.departure_time_var, 5, [2]))

        self.departure_time_entry = ttk.Entry(self.root, textvariable=self.departure_time_var)
        self.departure_time_entry.grid(row=6, column=1, padx=10, pady=10, sticky="w")
        self.departure_time_entry.bind("<KeyRelease>", lambda event: self.move_cursor(event, self.departure_time_entry))

        # Поля ввода даты прибытия
        self.arrival_date_label = tk.Label(self.root, text="Дата прибытия (дд.мм.гггг):")
        self.arrival_date_label.grid(row=7, column=0, padx=10, pady=10, sticky="e")

        self.arrival_date_var = tk.StringVar()
        self.arrival_date_var.trace_add("write",
                                        lambda *args: self.validate_date_entry(self.arrival_date_var, 10, [2, 5]))

        self.arrival_date_entry = ttk.Entry(self.root, textvariable=self.arrival_date_var)
        self.arrival_date_entry.grid(row=7, column=1, padx=10, pady=10, sticky="w")
        self.arrival_date_entry.bind("<KeyRelease>", lambda event: self.move_cursor(event, self.arrival_date_entry))

        # Поля ввода времени прибытия
        self.arrival_time_label = tk.Label(self.root, text="Время прибытия (чч:мм):")
        self.arrival_time_label.grid(row=8, column=0, padx=10, pady=10, sticky="e")

        self.arrival_time_var = tk.StringVar()
        self.arrival_time_var.trace_add("write", lambda *args: self.validate_time_entry(self.arrival_time_var, 5, [2]))

        self.arrival_time_entry = ttk.Entry(self.root, textvariable=self.arrival_time_var)
        self.arrival_time_entry.grid(row=8, column=1, padx=10, pady=10, sticky="w")
        self.arrival_time_entry.bind("<KeyRelease>", lambda event: self.move_cursor(event, self.arrival_time_entry))

        self.print_button = tk.Button(self.root, text="Печать", command=self.print_flight_details)
        self.print_button.grid(row=9, column=2, columnspan=2, pady=10)

        # Кнопка "Выполнить действие"
        self.action_button = tk.Button(self.root, text="Выполнить действие", command=self.perform_action)
        self.action_button.grid(row=9, column=0, columnspan=2, pady=10)

        # Таблица с рейсами
        self.table_label = tk.Label(self.root, text="Таблица рейсов:")
        self.table_label.grid(row=1, column=2, padx=10, pady=10, sticky="w")

        # Более узкие колонки
        column_widths = [100, 100, 80, 80, 120, 120, 120, 120]
        column_names = ["Пункт отправления", "Пункт прибытия", "Стоимость", "Номер рейса", "Дата отправления",
                        "Время отправления", "Дата прибытия", "Время прибытия"]

        self.table = ttk.Treeview(self.root, columns=(
        "From", "To", "Cost", "FlightNumber", "DepartureDate", "DepartureTime", "ArrivalDate", "ArrivalTime"),
                                  show="headings", height=20)

        for col, width, name in zip(
                ("From", "To", "Cost", "FlightNumber", "DepartureDate", "DepartureTime", "ArrivalDate", "ArrivalTime"),
                column_widths, column_names):
            self.table.heading(col, text=name)
            self.table.column(col, width=width)

        self.table.grid(row=2, column=2, rowspan=7, padx=10, pady=10)
        self.table.bind('<ButtonRelease-1>', self.bind_table_selection)

        # Вызов метода on_change_mode при создании виджетов
        self.on_change_mode()

    def bind_table_selection(self, event):
        # Получаем выбранный элемент в таблице
        selected_item = self.table.selection()
        if selected_item:
            # Получаем данные о рейсе из таблицы
            flight_data = {}
            for col, value in zip(
                    ["From", "To", "Cost", "FlightNumber", "DepartureDate", "DepartureTime", "ArrivalDate",
                     "ArrivalTime"], self.table.item(selected_item)['values']):
                flight_data[col] = value

            # Заполняем поля ввода данными выбранного рейса
            self.fill_entry_fields(flight_data)
            # Устанавливаем флаг редактирования рейса
            self.editing_flight = True

    def fill_entry_fields(self, flight_data):
        # Заполняем поля ввода данными выбранного рейса
        self.from_entry.delete(0, "end")
        self.from_entry.insert(0, flight_data["From"])

        self.to_entry.delete(0, "end")
        self.to_entry.insert(0, flight_data["To"])

        self.cost_entry.delete(0, "end")
        self.cost_entry.insert(0, flight_data["Cost"])

        self.flight_number_entry.delete(0, "end")
        self.flight_number_entry.insert(0, flight_data["FlightNumber"])

        self.departure_date_entry.delete(0, "end")
        self.departure_date_entry.insert(0, flight_data["DepartureDate"])

        self.departure_time_entry.delete(0, "end")
        self.departure_time_entry.insert(0, flight_data["DepartureTime"])

        self.arrival_date_entry.delete(0, "end")
        self.arrival_date_entry.insert(0, flight_data["ArrivalDate"])

        self.arrival_time_entry.delete(0, "end")
        self.arrival_time_entry.insert(0, flight_data["ArrivalTime"])

    def on_change_mode(self, *_):
        mode = self.mode.get()
        if mode == "Add":
            self.show_add_widgets()
        elif mode == "Remove":
            self.show_remove_widgets()
        elif mode == "Search":
            self.show_search_widgets()

    def show_add_widgets(self):
        self.clear_entries()
        self.from_entry["state"] = "normal"
        self.from_label["state"] = "normal"
        self.to_entry["state"] = "normal"
        self.to_label["state"] = "normal"
        self.cost_label["state"] = "normal"
        self.cost_entry["state"] = "normal"
        self.flight_number_label["state"] = "normal"
        self.flight_number_entry["state"] = "normal"
        self.departure_date_label["state"] = "normal"
        self.departure_date_entry["state"] = "normal"
        self.departure_time_label["state"] = "normal"
        self.departure_time_entry["state"] = "normal"
        self.arrival_date_label["state"] = "normal"
        self.arrival_date_entry["state"] = "normal"
        self.arrival_time_label["state"] = "normal"
        self.arrival_time_entry["state"] = "normal"
        self.update_table()

    def show_remove_widgets(self):
        self.clear_entries()
        self.from_entry["state"] = "disabled"
        self.from_label["state"] = "disabled"
        self.to_entry["state"] = "disabled"
        self.to_label["state"] = "disabled"
        self.cost_label["state"] = "disabled"
        self.cost_entry["state"] = "disabled"
        self.flight_number_label["state"] = "normal"
        self.flight_number_entry["state"] = "normal"
        self.departure_date_label["state"] = "disabled"
        self.departure_date_entry["state"] = "disabled"
        self.departure_time_label["state"] = "disabled"
        self.departure_time_entry["state"] = "disabled"
        self.arrival_date_label["state"] = "disabled"
        self.arrival_date_entry["state"] = "disabled"
        self.arrival_time_label["state"] = "disabled"
        self.arrival_time_entry["state"] = "disabled"
        self.update_table()

    def show_search_widgets(self):
        self.clear_entries()
        self.from_entry["state"] = "normal"
        self.from_label["state"] = "normal"
        self.to_entry["state"] = "normal"
        self.to_label["state"] = "normal"
        self.cost_label["state"] = "disabled"
        self.cost_entry["state"] = "disabled"
        self.flight_number_label["state"] = "disabled"
        self.flight_number_entry["state"] = "disabled"
        self.departure_date_label["state"] = "disabled"
        self.departure_date_entry["state"] = "disabled"
        self.departure_time_label["state"] = "disabled"
        self.departure_time_entry["state"] = "disabled"
        self.arrival_date_label["state"] = "disabled"
        self.arrival_date_entry["state"] = "disabled"
        self.arrival_time_label["state"] = "disabled"
        self.arrival_time_entry["state"] = "disabled"

    def perform_action(self):
        if self.mode.get() == "Add":
            if self.editing_flight:  # Если редактируем рейс
                self.edit_flight()
            else:
                self.add_flight()
        elif self.mode.get() == "Remove":
            self.remove_flight()
        elif self.mode.get() == "Search":
            self.search_flight()

    def edit_flight(self):
        # Редактирование рейса работает так же, как и добавление, но с изменением существующего рейса
        # Мы удаляем старый рейс и добавляем новый с обновленными данными
        self.remove_flight()
        self.add_flight()

    def add_flight(self):
        from_location = self.from_entry.get()
        to_location = self.to_entry.get()
        cost = self.cost_entry.get()
        flight_number = self.flight_number_entry.get()
        departure_date = self.departure_date_entry.get()
        departure_time = self.departure_time_entry.get()
        arrival_date = self.arrival_date_entry.get()
        arrival_time = self.arrival_time_entry.get()

        if not self.is_valid_flight_number(flight_number) or not departure_date:
            return

        # Округление значений даты и времени
        departure_date = self.round_date(departure_date)
        departure_time = self.round_time(departure_time)
        arrival_date = self.round_date(arrival_date)
        arrival_time = self.round_time(arrival_time)

        flight_data = {
            'From': from_location,
            'To': to_location,
            'Cost': cost,
            'FlightNumber': flight_number,
            'DepartureDate': departure_date,
            'DepartureTime': departure_time,
            'ArrivalDate': arrival_date,
            'ArrivalTime': arrival_time,
        }

        self.save_flight_data(flight_data)
        self.update_table()
        messagebox.showinfo("Info", f"Добавлен рейс с номером {flight_number}")

    def round_date(self, date_str):
        # Округление значений даты
        day, month, year = map(int, date_str.split('.'))
        day = min(day, 31)
        month = min(month, 12)
        year = min(year, 9999)
        return f"{day:02d}.{month:02d}.{year:04d}"

    def round_time(self, time_str):
        # Округление значений времени
        hours, minutes = map(int, time_str.split(':'))
        hours = min(hours, 23)
        minutes = min(minutes, 59)
        return f"{hours:02d}:{minutes:02d}"

    def is_valid_flight_number(self, flight_number):
        flights_filename = f'{self.username}_flights.json'
        if os.path.exists(flights_filename):
            with open(flights_filename, 'r') as file:
                user_flights = json.load(file)
                for flight in user_flights:
                    if flight.get('FlightNumber') == flight_number:
                        messagebox.showerror("Error", "Номер данного рейса уже существует.")
                        return False
        return True

    def remove_flight(self):
        flight_number = self.flight_number_entry.get()

        flights_filename = f'{self.username}_flights.json'
        if os.path.exists(flights_filename):
            with open(flights_filename, 'r') as file:
                user_flights = json.load(file)
        else:
            user_flights = []

        for idx, flight in enumerate(user_flights):
            if flight.get('FlightNumber') == flight_number:
                removed_flight = user_flights.pop(idx)
                with open(flights_filename, 'w') as file:
                    json.dump(user_flights, file)
                self.update_table()
                messagebox.showinfo("Info",
                                    f"Удален рейс с номером {flight_number}")  # Показываем сообщение после удаления
                return

        messagebox.showinfo("Info",
                            "Нет совпадений для удаления.")  # Показываем сообщение, если не найдено совпадений

    def search_flight(self):
        from_location = self.from_entry.get()
        to_location = self.to_entry.get()

        flights_filename = f'{self.username}_flights.json'
        if os.path.exists(flights_filename):
            with open(flights_filename, 'r') as file:
                user_flights = json.load(file)
        else:
            user_flights = []
        matching_flights = [flight for flight in user_flights if flight.get('From') == from_location and flight.get('To') == to_location]

        self.display_search_results(matching_flights)
        num_flights = len(matching_flights)
        messagebox.showinfo("Info", f"Найдено {num_flights} рейсов, удовлетворяющих условиям")

    def display_search_results(self, flights):
        # Очистим таблицу перед отображением результатов поиска
        for row in self.table.get_children():
            self.table.delete(row)

        for flight in flights:
            self.table.insert("", "end", values=(flight['From'], flight['To'], flight['Cost'], flight['FlightNumber'], flight['DepartureDate'], flight['DepartureTime'], flight['ArrivalDate'], flight['ArrivalTime']))

    def clear_entries(self):
        self.from_entry.delete(0, "end")
        self.to_entry.delete(0, "end")
        self.cost_entry.delete(0, "end")
        self.flight_number_entry.delete(0, "end")
        self.departure_date_entry.delete(0, "end")
        self.departure_time_entry.delete(0, "end")
        self.arrival_date_entry.delete(0, "end")
        self.arrival_time_entry.delete(0, "end")

    def update_table(self):
        flights_filename = f'{self.username}_flights.json'
        if os.path.exists(flights_filename):
            with open(flights_filename, 'r') as file:
                user_flights = json.load(file)
        else:
            user_flights = []

        # Очистим таблицу перед обновлением
        for row in self.table.get_children():
            self.table.delete(row)

        for flight in user_flights:
            self.table.insert("", "end", values=(flight['From'], flight['To'], flight['Cost'], flight['FlightNumber'], flight['DepartureDate'], flight['DepartureTime'], flight['ArrivalDate'], flight['ArrivalTime']))

    def save_flight_data(self, flight_data):
        flights_filename = f'{self.username}_flights.json'
        if os.path.exists(flights_filename):
            with open(flights_filename, 'r') as file:
                user_flights = json.load(file)
        else:
            user_flights = []

        user_flights.append(flight_data)

        with open(flights_filename, 'w') as file:
            json.dump(user_flights, file)

    @staticmethod
    def validate_cost(new_value):
        try:
            if new_value:
                float_value = float(new_value)
                if float_value < 0:
                    return False
            return True
        except ValueError:
            return False

    @staticmethod
    def validate_flight_number(new_value):
        # Проверка на длину номера рейса и допустимые символы
        return len(new_value) <= 7

    @staticmethod
    def validate_date(new_value):
        # Простая проверка формата даты (только цифры)
        return new_value.isdigit() or not new_value

    @staticmethod
    def validate_time(new_value):
        # Простая проверка формата времени (только цифры)
        return new_value.isdigit() or not new_value

    def validate_date_entry(self, var, max_length, dot_positions):
        # Автоматическое добавление точек при вводе даты и ограничение по количеству символов
        value = var.get()

        formatted_date = ''
        for i, char in enumerate(value):
            if char.isdigit():
                formatted_date += char
                if i + 1 in dot_positions and i + 1 != len(value):
                    formatted_date += '.'

        formatted_date = formatted_date[:max_length]
        var.set(formatted_date)

    def validate_time_entry(self, var, max_length, colon_positions):
        # Автоматическое добавление двоеточия при вводе времени и ограничение по количеству символов
        value = var.get()
        if value.isdigit() and len(value) in colon_positions and ':' not in value:
            var.set(value + ':')
        elif value and not value[-1].isdigit() or len(value) > max_length:
            var.set(value[:-1])

    def move_cursor(self, event, entry_widget):
        entry_widget.icursor("end")

    @staticmethod
    def alphabetic_validator(new_value):
        # Проверка, что вводимые символы являются буквами
        return True if new_value.isalpha() or not new_value else False

    def create_document(self, flight_data):
        document = Document()

        # Добавление заголовка
        document.add_heading('Детали рейса', 0)

        # Словарь с соответствиями имен столбцов
        column_mapping = {
            "From": "Пункт отправления",
            "To": "Пункт прибытия",
            "Cost": "Стоимость",
            "FlightNumber": "Номер рейса",
            "DepartureDate": "Дата отправления",
            "DepartureTime": "Время отправления",
            "ArrivalDate": "Дата прибытия",
            "ArrivalTime": "Время прибытия"
        }

        # Добавление данных о рейсе в таблицу
        table = document.add_table(rows=2, cols=8)
        table.style = 'Table Grid'
        table.autofit = False

        # Заголовки столбцов на русском
        for col, name in zip(table.columns, column_mapping.values()):
            col.cells[0].text = name

        # Данные о рейсе
        row = table.rows[1]
        for key, value in flight_data.items():
            row.cells[list(column_mapping.keys()).index(key)].text = str(value)

        # Добавление времени создания таблицы
        creation_time = datetime.datetime.now().strftime("%d.%m.%Y %H:%M")
        document.add_paragraph(f'Выписано {creation_time}')

        return document

    def print_flight_details(self):
        if self.mode.get() == "Add" or self.editing_flight:
            # Если добавляем новый рейс или редактируем существующий, то просто печатаем данные из полей ввода
            flight_data = {
                'From': self.from_entry.get(),
                'To': self.to_entry.get(),
                'Cost': self.cost_entry.get(),
                'FlightNumber': self.flight_number_entry.get(),
                'DepartureDate': self.departure_date_entry.get(),
                'DepartureTime': self.departure_time_entry.get(),
                'ArrivalDate': self.arrival_date_entry.get(),
                'ArrivalTime': self.arrival_time_entry.get(),
            }
        else:
            # Иначе, печатаем данные выбранного рейса в таблице
            selected_item = self.table.selection()
            if not selected_item:
                messagebox.showinfo("Уведомление", "Пожалуйста, выберите рейс.")
                return

            flight_data = {}
            for col, value in zip(["From", "To", "Cost", "FlightNumber", "DepartureDate", "DepartureTime", "ArrivalDate", "ArrivalTime"], self.table.item(selected_item)['values']):
                flight_data[col] = value

        document = self.create_document(flight_data)

        # Сохраняем документ
        file_path = f"FlightDetails_{flight_data['FlightNumber']}.docx"
        document.save(file_path)

        # Открываем документ в Word
        os.system(f'start {file_path}')


def register():
    user_data = load_user_data()
    username = entry_username.get()
    password = entry_password.get()

    # Проверяем, что введенные данные не пусты
    if not username or not password:
        messagebox.showerror("Ошибка", "Введите имя пользователя и пароль.")
        return

    if username in user_data:
        messagebox.showerror("Ошибка", "Пользователь уже существует.")
        return

    user_data[username] = hash_password(password)
    save_user_data(user_data)
    messagebox.showinfo("Успех", "Регистрация прошла успешно.")


# Функция входа
def login():
    user_data = load_user_data()
    username = entry_username.get()
    password = entry_password.get()
    if username in user_data and check_password(user_data[username], password):
        messagebox.showinfo("Успешно", "Вход в систему выполнен успешно.")
        root.withdraw()  # Скрыть окно входа
        flight_scheduler_app = FlightSchedulerApp(tk.Toplevel(), username)
    else:
        messagebox.showerror("Ошибка", "Неверное имя пользователя или пароль.")


root = tk.Tk()
root.geometry("300x200")
root.title("Авиакомпания")

label_username = tk.Label(root, text="Имя пользователя:")
label_username.pack()

# Создаем переменную для хранения значения в поле ввода логина
username_var = tk.StringVar()
# Ограничиваем максимальную длину в 32 символа
validate_username = root.register(validate_length)
entry_username = tk.Entry(root, textvariable=username_var, validate="key", validatecommand=(validate_username, "%P"))
entry_username.pack()

label_password = tk.Label(root, text="Пароль:")
label_password.pack()

# Создаем переменную для хранения значения в поле ввода пароля
password_var = tk.StringVar()
# Ограничиваем максимальную длину в 32 символа
validate_password = root.register(validate_length)
entry_password = tk.Entry(root, textvariable=password_var, show="*", validate="key",
                          validatecommand=(validate_password, "%P"))
entry_password.pack()

button_register = tk.Button(root, text="Регистрация", command=register)
button_register.pack()

button_login = tk.Button(root, text="Вход", command=login)
button_login.pack()

# Центрирование окна
center_window(root, 300, 200)

root.mainloop()


